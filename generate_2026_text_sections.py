#!/usr/bin/env python3
"""生成2026版通信工程专业本科人才培养方案-文字部分（第三、四部分）"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---- 全局样式设置 ----
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)  # 小四
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 辅助函数：设置表格单元格字体
def set_cell_font(cell, text, font_name='宋体', size=Pt(12), bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

# 辅助函数：添加正文段落
def add_body_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)  # 约2个中文字符
    p.paragraph_format.line_spacing = Pt(26)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 辅助函数：添加标题
def add_title(text, level='h3'):
    if level == 'h3':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = Pt(26)
        run = p.add_run(text)
        run.font.name = '黑体'
        run.font.size = Pt(15)  # 小三
        run.bold = True
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    return p

# ==========================================
# 第三部分：毕业要求
# ==========================================
add_title('三、毕业要求')

add_body_para('毕业要求总体描述：思想政治合格，在规定的年限内修完本专业人才培养方案所有课程，同时达到学校相应要求方能毕业。')

add_body_para('本专业依据《工程教育认证标准（2024版）》通用标准，制定了覆盖以下11条内容的毕业要求，并分解为37个可衡量、可评价的二级指标点。')

# 11条毕业要求数据
# 格式：(编号, 名称, 一级指标内容, [(二级编号, 二级名称, 二级内容), ...])
graduation_requirements = [
    ('4.3.1', '工程知识',
     '能够将数学、自然科学、计算、工程基础和通信工程专业知识用于解决通信领域的复杂工程问题。',
     [
         ('1.1', '数学知识',
          '掌握高等数学、工程数学等数学基础理论和方法，能够将数学知识用于通信领域复杂工程问题的建模、求解和分析。'),
         ('1.2', '自然科学知识',
          '掌握大学物理等自然科学基础知识，能够运用物理原理理解通信系统中的电磁场、信号传输等物理机制。'),
         ('1.3', '计算思维与方法',
          '理解计算思维的内涵，掌握数值分析、数据分析、统计以及计算机和信息科学的基础知识和方法，能够结合算力、算法和数据等要素，运用计算思维分析和解决通信领域工程问题。'),
         ('1.4', '工程基础知识',
          '掌握电路分析、电子线路、数字逻辑等工程基础知识，能够将工程基础理论用于通信系统的分析和设计。'),
         ('1.5', '通信工程专业知识',
          '掌握信号与系统、现代通信原理、通信网理论、移动通信等通信工程专业知识，能够综合运用专业知识解决通信领域的复杂工程问题。'),
     ]),
    ('4.3.2', '问题分析',
     '能够应用数学、自然科学和工程科学的基本原理，识别、表达并通过文献研究分析通信领域复杂工程问题，综合考虑可持续发展的要求，以获得有效结论。',
     [
         ('2.1', '问题识别与表达',
          '能够运用数学、自然科学和工程科学的基本原理，识别和判断通信领域复杂工程问题的关键环节和参数，并运用专业知识和数学模型方法进行正确表达。'),
         ('2.2', '文献研究与方案探索',
          '能够运用通信工程专业知识，通过文献检索、资料查询和研究分析，寻求通信领域复杂工程问题的不同解决方案。'),
         ('2.3', '综合分析',
          '能够综合运用数学、物理及专业知识，借助文献研究，并综合考虑可持续发展的要求，分析通信领域复杂工程问题，获得有效结论。'),
     ]),
    ('4.3.3', '设计/开发解决方案',
     '能够针对通信领域复杂工程问题设计和开发解决方案，设计满足特定需求的通信系统、单元（部件）或工艺流程，体现创新性，并从健康、安全与环境、全生命周期成本与净零碳要求、法律与伦理、社会与文化等角度考虑可行性。',
     [
         ('3.1', '工程设计与开发方法',
          '掌握通信工程设计和产品开发全周期、全流程的基本设计/开发方法和技术，了解影响设计目标和技术方案的各种因素。'),
         ('3.2', '系统与部件设计',
          '能够设计满足功能需求和性能指标要求的通信系统、功能单元或部件，提出针对复杂工程问题的解决方案。'),
         ('3.3', '创新性体现',
          '能够针对通信领域复杂工程问题，在方案设计环节中体现创新性，包括新方法、新架构或新技术的合理应用。'),
         ('3.4', '多因素约束下的可行性评估',
          '能够在设计环节综合考虑健康、安全与环境、全生命周期成本、净零碳要求、法律与伦理、社会与文化等因素，论证设计方案的可行性。'),
     ]),
    ('4.3.4', '研究',
     '能够基于科学原理并采用科学方法对通信领域复杂工程问题进行研究，包括设计实验、分析与解释数据，并通过信息综合得到合理有效的结论。',
     [
         ('4.1', '调研与方案设计',
          '能够基于科学原理，通过文献研究或相关方法，调研和分析通信系统中复杂工程问题，选择研究路线，设计合理的实验方案。'),
         ('4.2', '实验实施与数据采集',
          '能够根据实验方案构建通信实验系统，安全地开展实验，正确地采集实验数据。'),
         ('4.3', '数据分析与结论综合',
          '能够对实验结果进行分析和解释，并通过信息综合得到合理有效的结论。'),
     ]),
    ('4.3.5', '使用现代工具',
     '能够针对通信领域复杂工程问题，开发、选择与使用恰当的技术、资源、现代工程工具和信息技术工具，包括对复杂工程问题的预测与模拟，并能够理解其局限性。',
     [
         ('5.1', '工具选择与使用',
          '了解通信专业常用的仪器、设备、开发环境、仿真平台及编程工具的基本原理和使用方法，能够根据具体问题合理选择和使用适当的工具。'),
         ('5.2', '预测与模拟',
          '能够运用现代工具和网络资源，对通信领域复杂工程问题进行建模、仿真、预测与模拟，并分析结果的有效性。'),
         ('5.3', '工具局限性的理解',
          '能够理解现代工程工具和信息技术工具的局限性，在工程实践中合理评估工具适用性。'),
     ]),
    ('4.3.6', '工程与可持续发展',
     '在解决通信领域复杂工程问题时，能够基于工程相关背景知识，分析和评价工程实践对健康、安全、环境、法律以及经济和社会可持续发展的影响，并理解应承担的责任。',
     [
         ('6.1', '工程影响分析',
          '了解通信产业相关的技术标准、知识产权、行业要求和法律法规，能够分析和评价通信工程实践活动对社会、健康、安全、法律的影响。'),
         ('6.2', '可持续发展评价',
          '理解环境保护和可持续发展的理念和内涵，能够评价通信工程实践的生态环境影响和经济社会可持续发展影响。'),
         ('6.3', '责任意识',
          '理解通信工程师对公众健康、安全、环境以及社会可持续发展应承担的责任，能够在工程实践中自觉履行。'),
     ]),
    ('4.3.7', '工程伦理和职业规范',
     '有工程报国、为民造福的意识，具有人文社会科学素养和社会责任感，能够理解和践行工程伦理，在通信工程实践中遵守工程职业道德、规范和相关法律，履行责任。',
     [
         ('7.1', '工程报国与为民造福',
          '具有工程报国、为民造福的意识，树立正确的世界观、人生观和价值观，正确认识个体性与社会性的统一关系。'),
         ('7.2', '工程伦理理解与践行',
          '能够理解工程伦理的基本原则和内涵，在通信工程实践中自觉践行工程伦理，识别和规避工程实践中的伦理风险。'),
         ('7.3', '职业道德与法律遵守',
          '具有人文社会科学素养和社会责任感，能够在通信工程实践中遵守工程职业道德、规范和相关法律，自觉履行对公众及环境保护的责任。'),
     ]),
    ('4.3.8', '个人与团队',
     '能够在多样化、多学科背景下的团队中承担个体、团队成员以及负责人的角色。',
     [
         ('8.1', '个体贡献与角色定位',
          '具有有效沟通、合作共事等团队合作所需的基本素养，能够在团队中独立承担工作任务。'),
         ('8.2', '多样化团队协作',
          '理解团队成员身份、文化背景和工作场景的多样性（如面对面、远程和分布式协作），能够在多样化、多学科背景下与团队成员合作开展工作。'),
         ('8.3', '组织领导能力',
          '具有一定的组织管理能力，能够在多学科团队中承担负责人角色，组织、协调和指挥团队开展工作。'),
     ]),
    ('4.3.9', '沟通',
     '能够就通信领域复杂工程问题与业界同行及社会公众进行有效沟通和交流，包括撰写报告和设计文稿、陈述发言、清晰表达或回应指令；能够在跨文化背景下进行沟通和交流，理解、尊重语言和文化差异。',
     [
         ('9.1', '书面表达与文档撰写',
          '能够就通信领域复杂工程问题撰写技术报告、设计文稿和研究论文，具备规范的技术文档写作能力。'),
         ('9.2', '口头表达与陈述交流',
          '能够就通信领域复杂工程问题进行陈述发言，清晰表达技术观点，回应同行及公众的质疑和指令。'),
         ('9.3', '跨文化沟通与差异理解',
          '具备跟踪通信领域国际发展趋势和研究热点的能力，能够在跨文化背景下进行有效的沟通和交流，理解、尊重语言和文化差异。'),
     ]),
    ('4.3.10', '项目管理',
     '理解并掌握与通信工程项目相关的管理原理与经济决策方法，并能够在多学科环境中应用。',
     [
         ('10.1', '工程项目管理原理',
          '掌握通信领域工程项目中涉及的管理知识和原理，了解工程与产品的成本构成。'),
         ('10.2', '工程经济决策方法',
          '理解并掌握与通信工程项目相关的经济决策方法，能够在工程实施过程中进行经济性分析与决策。'),
         ('10.3', '多学科环境应用',
          '能够在多学科环境下，在通信工程项目设计开发解决方案的过程中，运用工程管理与经济决策方法。'),
     ]),
    ('4.3.11', '终身学习',
     '具有自主学习、终身学习和批判性思维的意识和能力，能够理解广泛的技术变革对工程和社会的影响，适应新技术变革。',
     [
         ('11.1', '自主学习意识与能力',
          '能够认识到不断探索和学习的必要性，对自主学习和终身学习具有正确的认识，具备自主学习的方法和能力。'),
         ('11.2', '批判性思维',
          '具有批判性思维能力，能够对工程实践中的既有方案和技术路线进行理性审视和反思，提出改进建议。'),
         ('11.3', '技术变革理解',
          '能够理解人工智能、大数据等广泛的技术变革对通信工程领域和社会发展产生的深刻影响。'),
         ('11.4', '适应新技术变革',
          '具备适应新技术变革的能力，能够主动更新知识结构，适应社会和行业的快速发展变化。'),
     ]),
]

# ---- 创建毕业要求表格 ----
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 设置表格宽度
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'), '9000')
tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)

# 设置列宽
for i, width in enumerate([1000, 4000, 1000, 3000]):
    for row in table.rows:
        row.cells[i].width = Cm(width / 567)

# 表头
headers = ['毕业要求\n一级指标', '一级指标内容', '毕业要求\n二级指标', '二级指标内容']
header_row = table.rows[0]
for i, header in enumerate(headers):
    set_cell_font(header_row.cells[i], header, font_name='黑体', size=Pt(10.5), bold=True)
    header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 填充数据
for req_id, req_name, req_content, sub_indicators in graduation_requirements:
    for j, (sub_id, sub_name, sub_content) in enumerate(sub_indicators):
        row = table.add_row()
        # 第一列：一级指标 - 只在第一条二级指标时填写
        if j == 0:
            set_cell_font(row.cells[0], f'{req_id}\n{req_name}', font_name='黑体', size=Pt(9), bold=True)
            set_cell_font(row.cells[1], req_content, font_name='宋体', size=Pt(9))
        else:
            set_cell_font(row.cells[0], '', font_name='宋体', size=Pt(9))
            set_cell_font(row.cells[1], '', font_name='宋体', size=Pt(9))
        # 第二列：二级指标
        set_cell_font(row.cells[2], sub_id, font_name='宋体', size=Pt(9))
        set_cell_font(row.cells[3], f'{sub_name}：{sub_content}', font_name='宋体', size=Pt(9))
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 合并第一、二列的单元格（按毕业要求）
# 删除每行第一列和第二列中空的内容，保留第一个非空行
# 找出需要合并的范围
merge_ranges = []
current_start = 1  # 从第二行开始（第一行是表头）
for i, (req_id, req_name, req_content, sub_indicators) in enumerate(graduation_requirements):
    count = len(sub_indicators)
    merge_ranges.append((current_start, current_start + count - 1))
    current_start += count

for start, end in merge_ranges:
    if end > start:
        table.cell(start, 0).merge(table.cell(end, 0))
        table.cell(start, 1).merge(table.cell(end, 1))

# 设置合并后单元格的垂直对齐
for start, end in merge_ranges:
    for col in [0, 1]:
        cell = table.cell(start, col)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)

doc.add_paragraph()  # 空行

# ==========================================
# 第四部分：毕业要求对培养目标的支撑矩阵
# ==========================================
add_title('四、毕业要求对培养目标的支撑矩阵')

add_body_para('本专业依据《工程教育认证标准（2024版）》制定的11条毕业要求，全面支撑4项培养目标的达成。各毕业要求对培养目标的支撑关系如下表所示：')

# 培养目标简称
objectives_short = [
    '培养目标1\n专业知识与计算思维',
    '培养目标2\n工程设计与创新能力',
    '培养目标3\n团队协作与项目管理能力',
    '培养目标4\n工程伦理与终身发展',
]

# 支撑矩阵数据 (毕业要求名称 → [目标1, 目标2, 目标3, 目标4])
support_matrix = [
    ('4.3.1 工程知识',                ['√', '',  '',  '']),
    ('4.3.2 问题分析',                ['√', '√', '',  '']),
    ('4.3.3 设计/开发解决方案',        ['',  '√', '',  '√']),
    ('4.3.4 研究',                    ['√', '√', '',  '']),
    ('4.3.5 使用现代工具',             ['',  '√', '',  '']),
    ('4.3.6 工程与可持续发展',         ['',  '',  '',  '√']),
    ('4.3.7 工程伦理和职业规范',       ['',  '',  '',  '√']),
    ('4.3.8 个人与团队',              ['',  '',  '√', '']),
    ('4.3.9 沟通',                    ['',  '',  '√', '']),
    ('4.3.10 项目管理',               ['',  '',  '√', '']),
    ('4.3.11 终身学习',               ['',  '√', '',  '√']),
]

# 创建支撑矩阵表格
matrix_table = doc.add_table(rows=1, cols=5)
matrix_table.style = 'Table Grid'
matrix_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 表头
matrix_header = ['毕业要求'] + objectives_short
header_row = matrix_table.rows[0]
for i, header in enumerate(matrix_header):
    set_cell_font(header_row.cells[i], header, font_name='黑体', size=Pt(9), bold=True)
    header_row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 填充数据
for req_name, supports in support_matrix:
    row = matrix_table.add_row()
    set_cell_font(row.cells[0], req_name, font_name='宋体', size=Pt(9))
    for j, val in enumerate(supports):
        set_cell_font(row.cells[j+1], val, font_name='宋体', size=Pt(12))
        row.cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加注释
doc.add_paragraph()
add_body_para('注：表中"√"表示该毕业要求对该项培养目标的达成具有重要支撑作用。')

# ---- 保存文件 ----
output_path = '2026版通信工程专业本科人才培养方案-毕业要求部分（草稿）.docx'
doc.save(output_path)
print(f'文件已生成: {output_path}')
print(f'毕业要求: {len(graduation_requirements)}条')
total_subs = sum(len(subs) for _, _, _, subs in graduation_requirements)
print(f'二级指标点: {total_subs}个')
